#!/usr/bin/env python3
"""Script to generate test files."""

import os
import sys
import logging
from pathlib import Path

# Add the parent directory to the path so we can import the modules
sys.path.append(str(Path(__file__).parent.parent.parent))

from src.utils.open_ai import OpenAIClient
from src.utils.general_utils import save_to_temp_file
from docx import Document
import shutil

# Create test data directory if it doesn't exist
os.makedirs('tests/data', exist_ok=True)

# Generate a sample resume in DOCX format
def generate_sample_resume():
    doc = Document()
    doc.add_heading('John Doe', 0)
    
    # Add contact information
    contact_info = doc.add_paragraph()
    contact_info.add_run('Email: john.doe@example.com | Phone: (123) 456-7890 | LinkedIn: linkedin.com/in/johndoe | GitHub: github.com/johndoe')
    
    # Add professional summary
    doc.add_heading('Professional Summary', level=1)
    doc.add_paragraph('Experienced software engineer with expertise in Python, JavaScript, and cloud technologies. Passionate about building scalable applications and solving complex problems.')
    
    # Add work experience
    doc.add_heading('Work Experience', level=1)
    
    # Job 1
    job1 = doc.add_paragraph()
    job1.add_run('Senior Software Engineer').bold = True
    job1.add_run(' | ABC Tech | 2020-Present\n')
    doc.add_paragraph('• Developed and maintained cloud-based applications using Python and AWS', style='List Bullet')
    doc.add_paragraph('• Led a team of 5 engineers to deliver a major product feature that increased user engagement by 30%', style='List Bullet')
    doc.add_paragraph('• Implemented CI/CD pipelines that reduced deployment time by 50%', style='List Bullet')
    
    # Job 2
    job2 = doc.add_paragraph()
    job2.add_run('Software Engineer').bold = True
    job2.add_run(' | XYZ Solutions | 2017-2020\n')
    doc.add_paragraph('• Built RESTful APIs using Django and Flask', style='List Bullet')
    doc.add_paragraph('• Optimized database queries that improved application performance by 40%', style='List Bullet')
    doc.add_paragraph('• Collaborated with cross-functional teams to deliver projects on time', style='List Bullet')
    
    # Add education
    doc.add_heading('Education', level=1)
    education = doc.add_paragraph()
    education.add_run('Bachelor of Science in Computer Science').bold = True
    education.add_run(' | University of Technology | 2013-2017\n')
    doc.add_paragraph('• GPA: 3.8/4.0', style='List Bullet')
    doc.add_paragraph('• Relevant coursework: Data Structures, Algorithms, Database Systems, Web Development', style='List Bullet')
    
    # Add skills
    doc.add_heading('Skills', level=1)
    skills = doc.add_paragraph()
    skills.add_run('Programming Languages: ').bold = True
    skills.add_run('Python, JavaScript, Java, SQL\n')
    skills.add_run('Frameworks & Libraries: ').bold = True
    skills.add_run('Django, Flask, React, Node.js\n')
    skills.add_run('Tools & Technologies: ').bold = True
    skills.add_run('AWS, Docker, Kubernetes, Git, CI/CD\n')
    
    # Save the document
    doc.save('tests/data/sample_resume.docx')
    print("Sample resume created at tests/data/sample_resume.docx")

# Generate a sample accomplishments file
def generate_sample_accomplishments():
    accomplishments = """
Professional Experience:
- Implemented a machine learning model that predicted customer churn with 85% accuracy, resulting in a 20% reduction in customer attrition
- Designed and developed a microservices architecture that improved system scalability and reduced downtime by 60%
- Optimized API response times by 70% through caching strategies and database indexing
- Mentored junior developers, leading to improved code quality and team productivity
- Received "Engineer of the Year" award for exceptional contributions to the company's flagship product

Personal Projects:
- Created an open-source library for data visualization that has over 500 stars on GitHub
- Developed a personal finance tracking application using React and Firebase
- Built a natural language processing tool that summarizes articles with 90% accuracy
- Contributed to several open-source projects, including Django and React

Education:
- Graduated with honors (magna cum laude)
- Completed capstone project on "Applying Machine Learning to Predict Stock Market Trends"
- Participated in ACM programming competitions, ranking in the top 10% nationally
- Served as president of the Computer Science Student Association
"""
    
    with open('tests/data/sample_accomplishments.txt', 'w') as f:
        f.write(accomplishments)
    print("Sample accomplishments created at tests/data/sample_accomplishments.txt")

# Generate all sample files
def generate_all_sample_files():
    generate_sample_resume()
    generate_sample_accomplishments()
    print("All sample files generated successfully")

if __name__ == "__main__":
    generate_all_sample_files()