#!/usr/bin/env python3
"""Module for writing resume to DOCX format."""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import docx2txt

from src.utils.resume_details import ResumeDetails


def set_paragraph_format(paragraph, font_size=9, space_after=Pt(0), line_spacing=1, left_indent=0):
    """Set paragraph formatting."""
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_after = space_after
    paragraph.line_spacing = line_spacing
    paragraph_format.left_indent = Inches(left_indent)
    for run in paragraph.runs:
        run.font.size = Pt(font_size)


def adjust_page_margins(document):
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)


def set_default_styles(document):
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(9)


def add_name(document, resume_details):
    if resume_details.name:
        name_paragraph = document.add_paragraph()
        name_run = name_paragraph.add_run(resume_details.name)
        name_run.bold = True
        name_run.font.size = Pt(18)
        name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_format(name_paragraph, font_size=18, space_after=Pt(0))


def add_contact_info(document, resume_details):
    contact_info_items = [
        resume_details.phone_number,
        resume_details.email,
        resume_details.address
    ]
    contact_info_items = [item for item in contact_info_items if item]
    if contact_info_items:
        contact_info = " | ".join(contact_info_items)
        contact_paragraph = document.add_paragraph()
        contact_paragraph.add_run(contact_info)
        contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_format(contact_paragraph, font_size=9, space_after=Pt(0))


def add_links(document, resume_details):
    link_items = []
    if resume_details.linkedin:
        link_items.append(f"LinkedIn: {resume_details.linkedin}")
    if resume_details.github:
        link_items.append(f"GitHub: {resume_details.github}")
    if link_items:
        links = " | ".join(link_items)
        links_paragraph = document.add_paragraph()
        links_paragraph.add_run(links)
        links_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_format(links_paragraph, font_size=9, space_after=Pt(6))


def add_professional_summary(document, resume_details):
    if resume_details.professional_summary:
        document.add_heading('Professional Summary', level=2)
        heading = document.paragraphs[-1]
        set_paragraph_format(heading, font_size=11, space_after=Pt(0))
        summary_paragraph = document.add_paragraph(resume_details.professional_summary)
        set_paragraph_format(summary_paragraph, font_size=9, space_after=Pt(6), left_indent=0.5)


def add_section(
    document, heading_title, entries,
    title_key='title', place_key='place', date_key='date',
    description_key='description', include_place=True, include_date=True
):
    if entries:
        # Remove entries that have no meaningful data
        entries = [
            entry for entry in entries
            if any([
                getattr(entry, title_key, None),
                getattr(entry, place_key, None) if include_place else False,
                getattr(entry, date_key, None) if include_date else False,
                getattr(entry, description_key, None)
            ])
        ]
        if entries:
            document.add_heading(heading_title, level=2)
            heading = document.paragraphs[-1]
            set_paragraph_format(heading, font_size=11, space_after=Pt(0))
            for entry in entries:
                title = getattr(entry, title_key, '')
                place = getattr(entry, place_key, '') if include_place and place_key else ''
                date = getattr(entry, date_key, '') if include_date and date_key else ''
                description = getattr(entry, description_key, []) or []

                p = document.add_paragraph()
                main_text = title
                if place:
                    main_text += f" - {place}"
                run = p.add_run(main_text)
                run.bold = True
                run.font.size = Pt(10)
                if date:
                    run = p.add_run(f" | {date}")
                    run.italic = True
                    run.font.size = Pt(9)
                set_paragraph_format(p, font_size=9, space_after=Pt(0), left_indent=0.5)

                for resp in description:
                    if resp:  # Check if the description is not empty
                        bullet = document.add_paragraph(style='List Bullet')
                        bullet_run = bullet.add_run(resp)
                        bullet_run.font.size = Pt(9)
                        bullet.paragraph_format.space_after = Pt(0)
                        bullet.line_spacing = 1
                        bullet.paragraph_format.left_indent = Inches(1)


def add_skills_and_languages(document, resume_details):
    skills_list = resume_details.skills if isinstance(resume_details.skills, list) else []
    languages_list = resume_details.languages if isinstance(resume_details.languages, list) else []
    skills_list = [skill for skill in skills_list if skill]
    languages_list = [lang for lang in languages_list if lang]
    combined_list = skills_list + languages_list
    if combined_list:
        document.add_heading('Skills & Languages', level=2)
        heading = document.paragraphs[-1]
        set_paragraph_format(heading, font_size=11, space_after=Pt(0))
        combined_text = ' | '.join(combined_list)
        combined_paragraph = document.add_paragraph(combined_text)
        set_paragraph_format(combined_paragraph, font_size=9, space_after=Pt(6), left_indent=0.5)


def write_resume_to_docx(resume_details: ResumeDetails, filename='result/resume.docx'):
    document = Document()
    adjust_page_margins(document)
    set_default_styles(document)
    if resume_details.name:
        add_name(document, resume_details)
    if any([resume_details.phone_number, resume_details.email, resume_details.address]):
        add_contact_info(document, resume_details)
    if any([resume_details.linkedin, resume_details.github]):
        add_links(document, resume_details)
    if resume_details.professional_summary:
        add_professional_summary(document, resume_details)

    # Add Work Experience
    if resume_details.work_experience:
        add_section(
            document,
            heading_title='Work Experience',
            entries=resume_details.work_experience,
            title_key='title',
            place_key='place',
            date_key='date',
            description_key='description',
            include_place=True,
            include_date=True
        )

    # Add Personal Projects
    if resume_details.personal_projects:
        add_section(
            document,
            heading_title='Personal Projects',
            entries=resume_details.personal_projects,
            title_key='title',
            description_key='description',
            include_place=False,
            include_date=False
        )

    # Add Education
    if resume_details.education:
        add_section(
            document,
            heading_title='Education',
            entries=resume_details.education,
            title_key='title',
            place_key='place',
            date_key='date',
            description_key='description',
            include_place=True,
            include_date=True
        )

    if resume_details.skills or resume_details.languages:
        add_skills_and_languages(document, resume_details)
    document.save(filename)
    print(f"Resume saved as {filename}")


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from a DOCX file.
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        Extracted text from the DOCX file
    """
    return docx2txt.process(file_path)
