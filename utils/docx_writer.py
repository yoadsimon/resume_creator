from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from utils.resume_details import ResumeDetails


def set_paragraph_format(paragraph, font_size=9, space_after=Pt(0), line_spacing=1, left_indent=0):
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_after = space_after
    paragraph.line_spacing = line_spacing
    paragraph_format.left_indent = Inches(left_indent)
    for run in paragraph.runs:
        run.font.size = Pt(font_size)


def adjust_page_margins(document):
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)


def set_default_styles(document):
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(9)


def add_name(document, resume_details):
    name_paragraph = document.add_paragraph()
    name_run = name_paragraph.add_run(resume_details.name)
    name_run.bold = True
    name_run.font.size = Pt(18)
    name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(name_paragraph, font_size=18, space_after=Pt(0))


def add_contact_info(document, resume_details):
    contact_info = f"{resume_details.phone_number} | {resume_details.email} | {resume_details.address}"
    contact_paragraph = document.add_paragraph()
    contact_paragraph.add_run(contact_info)
    contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(contact_paragraph, font_size=9, space_after=Pt(0))


def add_links(document, resume_details):
    links = f"LinkedIn: {resume_details.linkedin} | GitHub: {resume_details.github}"
    links_paragraph = document.add_paragraph()
    links_paragraph.add_run(links)
    links_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(links_paragraph, font_size=9, space_after=Pt(6))


def add_professional_summary(document, resume_details):
    document.add_heading('Professional Summary', level=2)
    heading = document.paragraphs[-1]
    set_paragraph_format(heading, font_size=11, space_after=Pt(0))
    summary_paragraph = document.add_paragraph(resume_details.professional_summary)
    set_paragraph_format(summary_paragraph, font_size=9, space_after=Pt(6), left_indent=0.5)


def add_section(document, heading_title, entries,
                title_key='title', place_key='place', date_key='date', description_key='description',
                include_place=True, include_date=True):
    document.add_heading(heading_title, level=2)
    heading = document.paragraphs[-1]
    set_paragraph_format(heading, font_size=11, space_after=Pt(0))
    for entry in entries:
        title = getattr(entry, title_key, '')
        place = getattr(entry, place_key, '') if include_place and place_key else ''
        date = getattr(entry, date_key, '') if include_date and date_key else ''
        description = getattr(entry, description_key, []) or []

        p = document.add_paragraph()
        main_text = title
        if place:
            main_text += f" - {place}"
        run = p.add_run(main_text)
        run.bold = True
        run.font.size = Pt(10)
        if date:
            run = p.add_run(f" | {date}")
            run.italic = True
            run.font.size = Pt(9)
        set_paragraph_format(p, font_size=9, space_after=Pt(0), left_indent=0.5)

        for resp in description:
            bullet = document.add_paragraph(style='List Bullet')
            bullet_run = bullet.add_run(resp)
            bullet_run.font.size = Pt(9)
            bullet.paragraph_format.space_after = Pt(0)
            bullet.line_spacing = 1
            bullet.paragraph_format.left_indent = Inches(1)


def add_skills_and_languages(document, resume_details):
    document.add_heading('Skills & Languages', level=2)
    heading = document.paragraphs[-1]
    set_paragraph_format(heading, font_size=11, space_after=Pt(0))
    skills_list = resume_details.skills if isinstance(resume_details.skills, list) else [resume_details.skills]
    languages_list = resume_details.languages if isinstance(resume_details.languages, list) else [
        resume_details.languages]
    combined_list = skills_list + languages_list
    combined_text = ' | '.join(combined_list)
    combined_paragraph = document.add_paragraph(combined_text)
    set_paragraph_format(combined_paragraph, font_size=9, space_after=Pt(6), left_indent=0.5)


def write_resume_to_docx(resume_details: ResumeDetails, filename='result/resume.docx'):
    document = Document()
    adjust_page_margins(document)
    set_default_styles(document)
    add_name(document, resume_details)
    add_contact_info(document, resume_details)
    add_links(document, resume_details)
    add_professional_summary(document, resume_details)

    # Add Work Experience
    add_section(
        document,
        heading_title='Work Experience',
        entries=resume_details.work_experience,
        title_key='title',
        place_key='place',
        date_key='date',
        description_key='description',
        include_place=True,
        include_date=True
    )

    # Add Personal Projects
    add_section(
        document,
        heading_title='Personal Projects',
        entries=resume_details.personal_projects,
        title_key='title',
        description_key='description',
        include_place=False,
        include_date=False
    )

    # Add Education
    add_section(
        document,
        heading_title='Education',
        entries=resume_details.education,
        title_key='title',
        place_key='place',
        date_key='date',
        description_key='description',
        include_place=True,
        include_date=True
    )

    add_skills_and_languages(document, resume_details)
    document.save(filename)
    print(f"Resume saved as {filename}")
